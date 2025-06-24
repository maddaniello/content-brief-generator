def generate_content_brief(self, data: Dict, keyword_analysis: Dict) -> str:
        """Genera il content brief usando OpenAI con analisi SEO avanzata e insights competitor"""
        
        # Analizza competitor in profonditàimport streamlit as st
import openai
import requests
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET
from urllib.parse import urljoin, urlparse
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import time
import json
from typing import List, Dict, Optional

# Configurazione della pagina
st.set_page_config(
    page_title="Content Brief Generator Pro",
    page_icon="📝",
    layout="wide"
)

# CSS personalizzato
st.markdown("""
<style>
    .main-header {
        text-align: center;
        color: #1f1f1f;
        margin-bottom: 2rem;
    }
    .section-header {
        color: #2c3e50;
        border-bottom: 2px solid #3498db;
        padding-bottom: 0.5rem;
        margin-top: 2rem;
    }
    .info-box {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #3498db;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #ffc107;
        margin: 1rem 0;
    }
    .error-box {
        background-color: #f8d7da;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #dc3545;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #28a745;
        margin: 1rem 0;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 0.5rem;
        color: white;
        text-align: center;
        margin: 0.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

class SEODataEnhancer:
    def __init__(self, semrush_api_key: str = None, serper_api_key: str = None):
        self.semrush_api_key = semrush_api_key
        self.serper_api_key = serper_api_key
        
    def get_semrush_keyword_data(self, keyword: str, country: str = "IT") -> Dict:
        """Ottiene dati dalle API di SEMrush per una keyword"""
        if not self.semrush_api_key:
            return {'status': 'error', 'message': 'SEMrush API key non configurata'}
        
        try:
            # API SEMrush per keyword overview
            url = "https://api.semrush.com/"
            params = {
                'type': 'phrase_organic',
                'key': self.semrush_api_key,
                'phrase': keyword,
                'database': country.lower(),
                'export_columns': 'Ph,Nq,Cp,Co,Nr,Td'
            }
            
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            
            # Parsing dei risultati SEMrush
            lines = response.text.strip().split('\n')
            if len(lines) > 1:
                data = lines[1].split(';')
                return {
                    'status': 'success',
                    'keyword': data[0] if len(data) > 0 else keyword,
                    'search_volume': int(data[1]) if len(data) > 1 and data[1].isdigit() else 0,
                    'cpc': float(data[2]) if len(data) > 2 and data[2].replace('.', '').isdigit() else 0,
                    'competition': float(data[3]) if len(data) > 3 and data[3].replace('.', '').isdigit() else 0,
                    'results_count': int(data[4]) if len(data) > 4 and data[4].isdigit() else 0,
                    'trend': data[5] if len(data) > 5 else ''
                }
            else:
                return {'status': 'no_data', 'keyword': keyword}
                
        except Exception as e:
            return {'status': 'error', 'keyword': keyword, 'error': str(e)}
    
    def get_semrush_related_keywords(self, keyword: str, country: str = "IT", limit: int = 50) -> List[Dict]:
        """Ottiene keyword correlate da SEMrush - aumentato limite per analisi più profonda"""
        if not self.semrush_api_key:
            return []
        
        try:
            url = "https://api.semrush.com/"
            params = {
                'type': 'phrase_related',
                'key': self.semrush_api_key,
                'phrase': keyword,
                'database': country.lower(),
                'export_columns': 'Ph,Nq,Cp,Co',
                'display_limit': limit
            }
            
            response = requests.get(url, params=params, timeout=10)
            response.raise_for_status()
            
            lines = response.text.strip().split('\n')
            related_keywords = []
            
            for line in lines[1:]:  # Skip header
                data = line.split(';')
                if len(data) >= 4:
                    related_keywords.append({
                        'keyword': data[0],
                        'search_volume': int(data[1]) if data[1].isdigit() else 0,
                        'cpc': float(data[2]) if data[2].replace('.', '').isdigit() else 0,
                        'competition': float(data[3]) if data[3].replace('.', '').isdigit() else 0
                    })
            
            return related_keywords
            
        except Exception as e:
            st.warning(f"Errore nell'ottenimento keyword correlate SEMrush: {str(e)}")
            return []
    
    def analyze_keyword_intent_patterns(self, related_keywords: List[Dict]) -> Dict:
        """Analizza i pattern di intento nelle keyword correlate"""
        intent_categories = {
            'informational': [],
            'navigational': [],
            'transactional': [],
            'commercial': []
        }
        
        # Pattern per classificare intent
        informational_patterns = ['come', 'cosa', 'quando', 'dove', 'perché', 'guida', 'tutorial', 'cos è', 'significa']
        transactional_patterns = ['acquista', 'compra', 'prezzo', 'costo', 'offerta', 'sconto', 'migliore', 'recensioni']
        commercial_patterns = ['confronto', 'vs', 'alternative', 'migliori', 'top', 'classifica', 'recensione']
        navigational_patterns = ['sito', 'ufficiale', 'login', 'accesso', 'brand']
        
        for kw_data in related_keywords:
            keyword = kw_data['keyword'].lower()
            
            if any(pattern in keyword for pattern in informational_patterns):
                intent_categories['informational'].append(kw_data)
            elif any(pattern in keyword for pattern in transactional_patterns):
                intent_categories['transactional'].append(kw_data)
            elif any(pattern in keyword for pattern in commercial_patterns):
                intent_categories['commercial'].append(kw_data)
            elif any(pattern in keyword for pattern in navigational_patterns):
                intent_categories['navigational'].append(kw_data)
            else:
                # Assegna alla categoria più probabile basata su CPC e competition
                if kw_data['cpc'] > 1.0 and kw_data['competition'] > 0.5:
                    intent_categories['transactional'].append(kw_data)
                elif kw_data['competition'] > 0.3:
                    intent_categories['commercial'].append(kw_data)
                else:
                    intent_categories['informational'].append(kw_data)
        
        return intent_categories
    
    def extract_topic_clusters(self, related_keywords: List[Dict]) -> Dict:
        """Estrae cluster tematici dalle keyword correlate"""
        clusters = {}
        
        for kw_data in related_keywords:
            keyword = kw_data['keyword'].lower()
            words = keyword.split()
            
            # Identifica temi principali
            for word in words:
                if len(word) > 3 and word not in ['come', 'cosa', 'quando', 'dove', 'perché', 'migliore', 'migliori']:
                    if word not in clusters:
                        clusters[word] = []
                    clusters[word].append(kw_data)
        
        # Filtra cluster con almeno 2 keyword
        filtered_clusters = {k: v for k, v in clusters.items() if len(v) >= 2}
        
        # Ordina per volume totale
        for cluster_name, keywords in filtered_clusters.items():
            total_volume = sum(kw['search_volume'] for kw in keywords)
            filtered_clusters[cluster_name] = {
                'keywords': keywords,
                'total_volume': total_volume,
                'avg_competition': sum(kw['competition'] for kw in keywords) / len(keywords)
            }
        
        return filtered_clusters
    
    def get_serper_search_data(self, query: str, country: str = "it") -> Dict:
        """Ottiene dati SERP da Serper API con analisi avanzata"""
        if not self.serper_api_key:
            return {'status': 'error', 'message': 'Serper API key non configurata'}
        
        try:
            url = "https://google.serper.dev/search"
            payload = {
                'q': query,
                'gl': country,
                'hl': 'it',
                'num': 10
            }
            headers = {
                'X-API-KEY': self.serper_api_key,
                'Content-Type': 'application/json'
            }
            
            response = requests.post(url, json=payload, headers=headers, timeout=10)
            response.raise_for_status()
            
            data = response.json()
            
            # Estrae People Also Ask con analisi dell'intent
            people_also_ask = []
            paa_intents = {'informational': [], 'commercial': [], 'transactional': []}
            
            if 'peopleAlsoAsk' in data:
                for paa in data['peopleAlsoAsk']:
                    question = paa.get('question', '')
                    people_also_ask.append(question)
                    
                    # Classifica intent della domanda
                    q_lower = question.lower()
                    if any(word in q_lower for word in ['come', 'cosa', 'quando', 'dove', 'perché']):
                        paa_intents['informational'].append(question)
                    elif any(word in q_lower for word in ['migliore', 'confronto', 'differenza', 'vs']):
                        paa_intents['commercial'].append(question)
                    elif any(word in q_lower for word in ['prezzo', 'costo', 'acquista', 'dove comprare']):
                        paa_intents['transactional'].append(question)
                    else:
                        paa_intents['informational'].append(question)
            
            # Estrae Related Searches
            related_searches = []
            if 'relatedSearches' in data:
                for rs in data['relatedSearches']:
                    related_searches.append(rs.get('query', ''))
            
            # Analizza Featured Snippet per pattern di successo
            featured_snippet = None
            snippet_analysis = {}
            if 'answerBox' in data:
                snippet_text = data['answerBox'].get('snippet', '')
                featured_snippet = {
                    'snippet': snippet_text,
                    'title': data['answerBox'].get('title', ''),
                    'link': data['answerBox'].get('link', '')
                }
                
                # Analizza struttura del snippet
                snippet_analysis = {
                    'word_count': len(snippet_text.split()),
                    'has_list': '•' in snippet_text or '-' in snippet_text or any(char.isdigit() and '.' in snippet_text for char in snippet_text),
                    'has_numbers': any(char.isdigit() for char in snippet_text),
                    'structure_type': 'list' if ('•' in snippet_text or '-' in snippet_text) else 'paragraph',
                    'starts_with_definition': snippet_text.lower().startswith(('è', 'sono', 'il', 'la', 'lo', 'una', 'un'))
                }
            
            # Estrae top 10 risultati con analisi dei domini
            organic_results = []
            domain_analysis = {}
            
            if 'organic' in data:
                for result in data['organic'][:10]:
                    domain = urlparse(result.get('link', '')).netloc
                    
                    organic_results.append({
                        'position': result.get('position', 0),
                        'title': result.get('title', ''),
                        'link': result.get('link', ''),
                        'snippet': result.get('snippet', ''),
                        'domain': domain
                    })
                    
                    # Conta frequenza domini
                    if domain:
                        domain_analysis[domain] = domain_analysis.get(domain, 0) + 1
            
            return {
                'status': 'success',
                'query': query,
                'people_also_ask': people_also_ask,
                'paa_intents': paa_intents,
                'related_searches': related_searches,
                'featured_snippet': featured_snippet,
                'snippet_analysis': snippet_analysis,
                'organic_results': organic_results,
                'domain_analysis': domain_analysis,
                'total_results': data.get('searchInformation', {}).get('totalResults', 0)
            }
            
        except Exception as e:
            return {'status': 'error', 'query': query, 'error': str(e)}

class ContentBriefGenerator:
    def __init__(self, api_key: str, seo_enhancer: SEODataEnhancer = None):
        self.client = openai.OpenAI(api_key=api_key)
        self.seo_enhancer = seo_enhancer or SEODataEnhancer()
    
    def analyze_competitor_content(self, competitors: List[Dict]) -> Dict:
        """Analizza in profondità il contenuto dei competitor per estrarre insight strategici"""
        analysis = {
            'common_topics': {},
            'content_gaps': [],
            'structural_patterns': {},
            'keyword_usage_patterns': {},
            'content_depth_analysis': {},
            'unique_angles': {}
        }
        
        # Analizza ogni competitor
        for comp in competitors:
            content = comp['content'].lower()
            
            # Estrae topic principali
            topic_keywords = []
            words = content.split()
            
            # Pattern per identificare topic importanti
            for i, word in enumerate(words):
                if len(word) > 4:  # Parole significative
                    # Context window per capire importanza
                    context_before = words[max(0, i-3):i]
                    context_after = words[i+1:min(len(words), i+4)]
                    
                    # Se la parola appare in contesti importanti
                    if any(ctx in ['importante', 'fondamentale', 'essenziale', 'principale', 'primo', 'migliore'] for ctx in context_before + context_after):
                        topic_keywords.append(word)
            
            # Conta frequenza topic
            for topic in topic_keywords:
                if topic not in analysis['common_topics']:
                    analysis['common_topics'][topic] = 0
                analysis['common_topics'][topic] += 1
            
            # Analizza struttura contenuto
            headings = comp.get('headings', '').split('\n')
            heading_patterns = []
            for heading in headings:
                if heading.strip():
                    h_level = heading.split(':')[0] if ':' in heading else 'H1'
                    h_text = heading.split(':', 1)[1].strip() if ':' in heading else heading
                    
                    # Pattern strutturali
                    if 'come' in h_text.lower():
                        heading_patterns.append('how_to')
                    elif 'cosa' in h_text.lower():
                        heading_patterns.append('what_is')
                    elif 'perché' in h_text.lower():
                        heading_patterns.append('why')
                    elif 'migliori' in h_text.lower() or 'top' in h_text.lower():
                        heading_patterns.append('best_list')
                    elif 'confronto' in h_text.lower() or 'vs' in h_text.lower():
                        heading_patterns.append('comparison')
            
            # Accumula pattern strutturali
            for pattern in heading_patterns:
                if pattern not in analysis['structural_patterns']:
                    analysis['structural_patterns'][pattern] = 0
                analysis['structural_patterns'][pattern] += 1
            
            # Analizza profondità contenuto
            paragraphs = comp.get('paragraphs', [])
            analysis['content_depth_analysis'][f"competitor_{comp['competitor_number']}"] = {
                'total_paragraphs': len(paragraphs),
                'avg_paragraph_length': sum(len(p.split()) for p in paragraphs) / len(paragraphs) if paragraphs else 0,
                'word_count': comp['word_count'],
                'has_lists': 'lista' in content or 'elenco' in content or '•' in comp['content'],
                'has_examples': 'esempio' in content or 'ad esempio' in content,
                'technical_depth': content.count('tecnic') + content.count('specific') + content.count('dettagli')
            }
        
        # Identifica gap di contenuto
        all_topics = set()
        for comp in competitors:
            words = comp['content'].lower().split()
            comp_topics = [w for w in words if len(w) > 5]
            all_topics.update(comp_topics)
        
        # Topic coverage analysis
        topic_coverage = {}
        for topic in all_topics:
            coverage_count = sum(1 for comp in competitors if topic in comp['content'].lower())
            topic_coverage[topic] = coverage_count
        
        # Identifica topic coperti da pochi competitor (opportunità)
        analysis['content_gaps'] = [
            topic for topic, count in topic_coverage.items() 
            if count == 1 and len(topic) > 6  # Topic unici e significativi
        ][:10]
        
        return analysis
    
    def extract_search_intent_insights(self, keyword_analysis: Dict) -> Dict:
        """Estrae insight avanzati sull'intento di ricerca"""
        insights = {
            'primary_intent': 'informational',
            'intent_distribution': {},
            'content_suggestions': {},
            'user_journey_stage': 'awareness',
            'competition_level': 'medium'
        }
        
        # Analizza dati SEMrush
        if keyword_analysis.get('semrush_data', {}).get('status') == 'success':
            semrush = keyword_analysis['semrush_data']
            
            # Determina intent principale da CPC e competition
            cpc = semrush.get('cpc', 0)
            competition = semrush.get('competition', 0)
            
            if cpc > 2.0 and competition > 0.7:
                insights['primary_intent'] = 'transactional'
                insights['user_journey_stage'] = 'decision'
            elif cpc > 1.0 and competition > 0.4:
                insights['primary_intent'] = 'commercial'
                insights['user_journey_stage'] = 'consideration'
            else:
                insights['primary_intent'] = 'informational'
                insights['user_journey_stage'] = 'awareness'
            
            # Livello competizione
            if competition > 0.8:
                insights['competition_level'] = 'high'
            elif competition > 0.4:
                insights['competition_level'] = 'medium'
            else:
                insights['competition_level'] = 'low'
        
        # Analizza keyword correlate
        if keyword_analysis.get('related_keywords'):
            intent_patterns = keyword_analysis.get('intent_categories', {})
            total_keywords = len(keyword_analysis['related_keywords'])
            
            if total_keywords > 0:
                for intent, keywords in intent_patterns.items():
                    insights['intent_distribution'][intent] = len(keywords) / total_keywords
        
        # Analizza PAA per suggerimenti contenuto
        if keyword_analysis.get('serper_data', {}).get('paa_intents'):
            paa_intents = keyword_analysis['serper_data']['paa_intents']
            
            insights['content_suggestions'] = {
                'faq_section_needed': len(paa_intents.get('informational', [])) > 2,
                'comparison_section_needed': len(paa_intents.get('commercial', [])) > 1,
                'pricing_section_needed': len(paa_intents.get('transactional', [])) > 1,
                'how_to_section_needed': any('come' in q.lower() for q in keyword_analysis['serper_data'].get('people_also_ask', []))
            }
        
        return insights
    
    def get_sitemap_urls(self, sitemap_url: str) -> List[str]:
        """Estrae le URL dalla sitemap con gestione di sitemap multiple"""
        urls = []
        processed_sitemaps = set()
        
        def process_sitemap(url: str):
            if url in processed_sitemaps:
                return
            processed_sitemaps.add(url)
            
            try:
                headers = {
                    'User-Agent': 'Mozilla/5.0 (compatible; SEO-Tool/1.0; +http://www.example.com/bot)'
                }
                response = requests.get(url, headers=headers, timeout=15)
                response.raise_for_status()
                
                root = ET.fromstring(response.content)
                namespaces = {
                    'ns': 'http://www.sitemaps.org/schemas/sitemap/0.9',
                    'sitemap': 'http://www.sitemaps.org/schemas/sitemap/0.9'
                }
                
                sitemapindex = root.findall('.//ns:sitemap', namespaces)
                if sitemapindex:
                    for sitemap in sitemapindex[:10]:
                        loc_elem = sitemap.find('ns:loc', namespaces)
                        if loc_elem is not None:
                            process_sitemap(loc_elem.text)
                else:
                    for url_elem in root.findall('.//ns:url', namespaces):
                        loc_elem = url_elem.find('ns:loc', namespaces)
                        if loc_elem is not None:
                            urls.append(loc_elem.text)
                            if len(urls) >= 100:
                                break
                
            except ET.ParseError:
                try:
                    response = requests.get(url, headers=headers, timeout=15)
                    text = response.text
                    url_pattern = r'https?://[^\s<>"\']+(?:/[^\s<>"\']*)?'
                    found_urls = re.findall(url_pattern, text)
                    urls.extend(found_urls[:50])
                except:
                    pass
                    
            except Exception as e:
                st.warning(f"Errore nell'elaborazione della sitemap {url}: {str(e)}")
        
        try:
            process_sitemap(sitemap_url)
            return list(set(urls))
        except Exception as e:
            st.error(f"Errore critico nell'estrazione della sitemap: {str(e)}")
            return []
    
    def analyze_keywords_with_apis(self, keywords: str) -> Dict:
        """Analizza le keyword usando SEMrush e Serper con analisi avanzata"""
        keyword_list = [k.strip() for k in keywords.split(',') if k.strip()]
        main_keyword = keyword_list[0] if keyword_list else ""
        
        # Dati SEMrush
        semrush_data = {}
        related_keywords = []
        intent_categories = {}
        topic_clusters = {}
        
        if main_keyword:
            semrush_data = self.seo_enhancer.get_semrush_keyword_data(main_keyword)
            if semrush_data.get('status') == 'success':
                related_keywords = self.seo_enhancer.get_semrush_related_keywords(main_keyword, limit=50)
                if related_keywords:
                    intent_categories = self.seo_enhancer.analyze_keyword_intent_patterns(related_keywords)
                    topic_clusters = self.seo_enhancer.extract_topic_clusters(related_keywords)
        
        # Dati Serper con analisi avanzata
        serper_data = {}
        if main_keyword:
            serper_data = self.seo_enhancer.get_serper_search_data(main_keyword)
        
        return {
            'main_keyword': main_keyword,
            'all_keywords': keyword_list,
            'semrush_data': semrush_data,
            'related_keywords': related_keywords,
            'intent_categories': intent_categories,
            'topic_clusters': topic_clusters,
            'serper_data': serper_data
        }
    
    def generate_content_brief(self, data: Dict, keyword_analysis: Dict) -> str:
        """Genera il content brief usando OpenAI con analisi SEO avanzata e insights competitor"""
        
        # Analizza competitor in profondità
        competitor_insights = self.analyze_competitor_content(data['competitors'])
        
        # Estrae insight sull'intento di ricerca
        search_intent_insights = self.extract_search_intent_insights(keyword_analysis)
        
        # Prepara i dati dei competitor con analisi avanzata
        competitor_data = ""
        for comp in data['competitors']:
            competitor_data += f"\n--- COMPETITOR {comp['competitor_number']} ---\n"
            competitor_data += f"URL: {comp['url']}\n"
            if comp['title'] != f"Competitor {comp['competitor_number']}":
                competitor_data += f"Titolo: {comp['title']}\n"
            if comp['meta_description']:
                competitor_data += f"Meta Description: {comp['meta_description']}\n"
            if comp['headings']:
                competitor_data += f"Struttura titoli identificata:\n{comp['headings']}\n"
            competitor_data += f"Numero parole: {comp['word_count']}\n"
            competitor_data += f"Contenuto completo: {comp['content'][:3000]}...\n"  # Primi 3000 caratteri
        
        # Prepara analisi competitor avanzata
        competitor_analysis = f"""
ANALISI AVANZATA COMPETITOR:
- Topic più comuni: {', '.join(list(competitor_insights['common_topics'].keys())[:10])}
- Pattern strutturali dominanti: {', '.join(competitor_insights['structural_patterns'].keys())}
- Gap di contenuto identificati: {', '.join(competitor_insights['content_gaps'][:5])}
- Profondità media contenuto: {sum(comp['word_count'] for comp in data['competitors']) / len(data['competitors']):.0f} parole
- Elementi strutturali mancanti: {', '.join([k for k, v in competitor_insights['content_depth_analysis'].items()])}
"""
        
        # Prepara dati SEMrush avanzati
        semrush_info = ""
        if keyword_analysis['semrush_data'].get('status') == 'success':
            sd = keyword_analysis['semrush_data']
            semrush_info = f"""
DATI SEMRUSH KEYWORD PRINCIPALE "{keyword_analysis['main_keyword']}":
- Volume di ricerca mensile: {sd.get('search_volume', 'N/A')}
- CPC: €{sd.get('cpc', 'N/A')} (Indicatore intent: {'Transactional' if sd.get('cpc', 0) > 2 else 'Commercial' if sd.get('cpc', 0) > 1 else 'Informational'})
- Competizione: {sd.get('competition', 'N/A')}/1.0 (Livello: {'Alto' if sd.get('competition', 0) > 0.7 else 'Medio' if sd.get('competition', 0) > 0.4 else 'Basso'})
- Risultati totali: {sd.get('results_count', 'N/A')}
"""
        
        # Prepara keyword correlate per intent
        related_kw_by_intent = ""
        if keyword_analysis.get('intent_categories'):
            for intent, keywords in keyword_analysis['intent_categories'].items():
                if keywords:
                    related_kw_by_intent += f"\nKEYWORD {intent.upper()}:\n"
                    top_keywords = sorted(keywords, key=lambda x: x['search_volume'], reverse=True)[:5]
                    for kw in top_keywords:
                        related_kw_by_intent += f"- {kw['keyword']} (Vol: {kw['search_volume']}, Comp: {kw['competition']:.2f})\n"
        
        # Prepara cluster tematici
        topic_clusters_info = ""
        if keyword_analysis.get('topic_clusters'):
            topic_clusters_info = "CLUSTER TEMATICI DA SEMRUSH:\n"
            sorted_clusters = sorted(keyword_analysis['topic_clusters'].items(), 
                                   key=lambda x: x[1]['total_volume'], reverse=True)[:5]
            for cluster_name, cluster_data in sorted_clusters:
                topic_clusters_info += f"- Tema '{cluster_name}': {cluster_data['total_volume']} vol. totale, {len(cluster_data['keywords'])} keyword\n"
        
        # Prepara dati Serper avanzati
        serper_info = ""
        if keyword_analysis['serper_data'].get('status') == 'success':
            sd = keyword_analysis['serper_data']
            
            if sd.get('people_also_ask'):
                serper_info += "PEOPLE ALSO ASK DA GOOGLE (CLASSIFICATE PER INTENT):\n"
                paa_intents = sd.get('paa_intents', {})
                for intent, questions in paa_intents.items():
                    if questions:
                        serper_info += f"\n{intent.upper()}:\n"
                        for q in questions[:3]:
                            serper_info += f"- {q}\n"
            
            if sd.get('related_searches'):
                serper_info += f"\nRICERCHE CORRELATE DA GOOGLE:\n"
                for rs in sd['related_searches'][:5]:
                    serper_info += f"- {rs}\n"
            
            if sd.get('featured_snippet'):
                snippet_analysis = sd.get('snippet_analysis', {})
                serper_info += f"\nFEATURED SNIPPET ATTUALE - ANALISI STRUTTURALE:\n"
                serper_info += f"Titolo: {sd['featured_snippet']['title']}\n"
                serper_info += f"Lunghezza: {snippet_analysis.get('word_count', 0)} parole\n"
                serper_info += f"Tipo struttura: {snippet_analysis.get('structure_type', 'paragraph')}\n"
                serper_info += f"Contiene liste: {'Sì' if snippet_analysis.get('has_list') else 'No'}\n"
                serper_info += f"Contiene numeri: {'Sì' if snippet_analysis.get('has_numbers') else 'No'}\n"
                serper_info += f"Snippet: {sd['featured_snippet']['snippet'][:200]}...\n"
        
        # Prepara insight intento di ricerca
        intent_insights = f"""
ANALISI INTENTO DI RICERCA AVANZATA:
- Intent principale: {search_intent_insights['primary_intent']}
- Fase user journey: {search_intent_insights['user_journey_stage']}
- Livello competizione: {search_intent_insights['competition_level']}
- Distribuzione intent: {search_intent_insights.get('intent_distribution', {})}
- Sezioni consigliate: {search_intent_insights.get('content_suggestions', {})}
"""
        
        # Prepara le URL interne
        internal_urls = "\n".join(data['sitemap_urls'][:30]) if data['sitemap_urls'] else data.get('manual_urls', 'Nessuna URL interna disponibile')
        
        prompt = f"""
Sei un esperto SEO content strategist e data analyst specializzato in E-E-A-T. Devi creare un content brief ESTREMAMENTE DETTAGLIATO e ACTIONABLE basato su DATI SEO REALI e ANALISI AVANZATA dei competitor.

INFORMAZIONI CLIENTE:
- Brand: {data['brand']}
- Sito web: {data['website']}
- Argomento: {data['topic']}
- Keyword principali: {data['keywords']}
- Domande frequenti inserite: {data['faqs']}
- Tone of voice: {', '.join(data['tone_of_voice'])}

{semrush_info}

{related_kw_by_intent}

{topic_clusters_info}

{serper_info}

{intent_insights}

{competitor_analysis}

ANALISI DETTAGLIATA COMPETITOR:
{competitor_data}

URL INTERNE DISPONIBILI (per link interni):
{internal_urls}

ISTRUZIONI SPECIFICHE:
- Il brand "{data['brand']}" DEVE apparire nel meta title alla fine
- Il brand "{data['brand']}" DEVE apparire nella meta description
- Usa la capitalizzazione naturale italiana (prima lettera maiuscola il resto minuscolo, ad esempio 'Mutuo per Acquisto Garage: Tutto Quello che Devi Sapere' NON va bene, 'Mutuo per acquisto garage: tutto quello che devi sapere' va bene
- Utilizza TUTTI i dati reali per creare suggerimenti specifici e actionable
- Per ogni H2/H3 fornisci istruzioni DETTAGLIATE su cosa scrivere all'interno di quel paragrafo tramite degli elenchi puntati dettagliati
- Sfrutta i gap dei competitor per opportunità uniche
- Integra keyword correlate per intent specifici
- Rispondi strategicamente alle PAA classificate per intent

Genera un content brief che includa:

1. **STRATEGIA SEO DATA-DRIVEN AVANZATA**
   - Analisi intent basata su CPC ({keyword_analysis['semrush_data'].get('cpc', 0)}€) e competition ({keyword_analysis['semrush_data'].get('competition', 0)})
   - Strategia per fase user journey: {search_intent_insights['user_journey_stage']}
   - Piano per battere featured snippet attuale (se presente)
   - Sfruttamento gap competitor identificati

2. **META OTTIMIZZATI CON DATI REALI**
   - Meta title (50-60 caratteri) ottimizzato per volume {keyword_analysis['semrush_data'].get('search_volume', 0)}
   - Meta description che incorpora PAA ad alto search intent
   - Keyword correlate strategiche da integrare

3. **STRUTTURA CONTENUTO ESTREMAMENTE DETTAGLIATA**
   
   **H1 OTTIMIZZATO:**
   - H1 specifico con keyword principale
   - Giustificazione scelta basata su dati competitor
   
   **INTRODUZIONE STRATEGICA:**
   - Cosa scrivere nei primi 2-3 paragrafi
   - Come incorporare keyword principale naturalmente
   - Hook basato su gap competitor identificati
   - Lunghezza ottimale e elementi da includere
   
   **SEZIONI H2 CON ISTRUZIONI DETTAGLIATE:**
   Per ogni H2 fornisci:
   - Titolo H2 ottimizzato per keyword correlate specifiche
   - 4-6 bullet point DETTAGLIATI su cosa scrivere in quel paragrafo
   - Keyword correlate specifiche da integrare (con volumi di ricerca)
   - PAA specifiche da rispondere in quella sezione
   - Esempi concreti da includere
   - Lunghezza paragrafo consigliata
   - Elementi aggiuntivi (liste, tabelle, immagini)
   - Link interni strategici con anchor text specifiche
   
   **SOTTOSEZIONI H3 QUANDO NECESSARIO:**
   - H3 per approfondimenti specifici
   - Istruzioni puntuali su contenuto
   - Keyword long-tail da targetizzare
   
   **SEZIONI SPECIALI BASATE SU DATI:**
   - Sezione FAQ (se PAA > 3)
   - Sezione confronto (se keyword commercial presenti)
   - Sezione how-to (se keyword informational dominanti)
   - Sezione pricing/costi (se CPC > 1€)

4. **STRATEGIA PEOPLE ALSO ASK AVANZATA**
   Per ogni PAA da Google:
   - In quale sezione H2/H3 rispondere
   - Come strutturare la risposta (lunghezza, formato)
   - Keyword correlate da includere nella risposta
   - Opportunità per featured snippet

5. **INTEGRAZIONE KEYWORD CORRELATE PER TOPIC CLUSTER**
   Per ogni cluster tematico identificato:
   - Dove integrare le keyword del cluster
   - Densità ottimale basata su competition
   - Long-tail opportunities ad alto volume

6. **STRATEGIA LINK INTERNI DATA-DRIVEN**
   - Link basati su keyword correlate e volumi
   - Anchor text ottimizzate per topic cluster
   - Distribuzione strategica per massimizzare ranking

7. **ELEMENTI E-E-A-T SPECIFICI**
   - Fonti autorevoli che competitor non usano
   - Dati statistici più recenti
   - Esempi pratici basati su ricerche correlate reali
   - Authority signals da includere

8. **PIANO IMPLEMENTAZIONE COPYWRITER**
   - Checklist step-by-step per copywriter
   - Metriche da raggiungere (lunghezza, keyword density)
   - Elementi obbligatori per ogni sezione
   - KPI di successo previsti

Crea un brief che permetta al copywriter di scrivere contenuto SUPERIORE ai competitor utilizzando ESCLUSIVAMENTE dati reali e analisi avanzate. Ogni suggerimento deve essere specifico, actionable e basato sui dati forniti.
"""

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "Sei un esperto SEO data analyst e content strategist che crea content brief estremamente dettagliati e actionable utilizzando dati reali di SEMrush, Serper e analisi competitor avanzate per garantire posizionamenti top su Google."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=4000,
                temperature=0.3  # Più deterministico per dati specifici
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            st.error(f"Errore nella generazione del content brief: {str(e)}")
            return "Errore nella generazione del contenuto."
        
        # Prepara dati SEMrush
        semrush_info = ""
        if keyword_analysis['semrush_data'].get('status') == 'success':
            sd = keyword_analysis['semrush_data']
            semrush_info = f"""
DATI SEMRUSH KEYWORD PRINCIPALE "{keyword_analysis['main_keyword']}":
- Volume di ricerca mensile: {sd.get('search_volume', 'N/A')}
- CPC: €{sd.get('cpc', 'N/A')}
- Competizione: {sd.get('competition', 'N/A')}/1.0
- Risultati totali: {sd.get('results_count', 'N/A')}
"""
        
        # Prepara keyword correlate
        related_kw = ""
        if keyword_analysis['related_keywords']:
            related_kw = "KEYWORD CORRELATE DA SEMRUSH:\n"
            for kw in keyword_analysis['related_keywords'][:10]:
                related_kw += f"- {kw['keyword']} (Vol: {kw['search_volume']}, Comp: {kw['competition']})\n"
        
        # Prepara dati Serper
        serper_info = ""
        if keyword_analysis['serper_data'].get('status') == 'success':
            sd = keyword_analysis['serper_data']
            
            if sd.get('people_also_ask'):
                serper_info += "PEOPLE ALSO ASK DA GOOGLE:\n"
                for paa in sd['people_also_ask'][:8]:
                    serper_info += f"- {paa}\n"
            
            if sd.get('related_searches'):
                serper_info += "\nRICERCHE CORRELATE DA GOOGLE:\n"
                for rs in sd['related_searches'][:5]:
                    serper_info += f"- {rs}\n"
            
            if sd.get('featured_snippet'):
                serper_info += f"\nFEATURED SNIPPET ATTUALE:\n"
                serper_info += f"Titolo: {sd['featured_snippet']['title']}\n"
                serper_info += f"Snippet: {sd['featured_snippet']['snippet'][:200]}...\n"
        
        # Prepara le URL interne
        internal_urls = "\n".join(data['sitemap_urls'][:30]) if data['sitemap_urls'] else data.get('manual_urls', 'Nessuna URL interna disponibile')
        
        prompt = f"""
Sei un esperto SEO copywriter e content strategist specializzato in E-E-A-T (Experience, Expertise, Authoritativeness, Trustworthiness). Devi creare un content brief dettagliato basato su DATI SEO REALI per un articolo che raggiunga il massimo punteggio E-E-A-T e batta la concorrenza.

INFORMAZIONI CLIENTE:
- Brand: {data['brand']}
- Sito web: {data['website']}
- Argomento: {data['topic']}
- Keyword principali: {data['keywords']}
- Domande frequenti inserite: {data['faqs']}
- Tone of voice: {', '.join(data['tone_of_voice'])}

{semrush_info}

{related_kw}

{serper_info}

ANALISI COMPETITOR:
{competitor_data}

URL INTERNE DISPONIBILI (per link interni):
{internal_urls}

IMPORTANTE: 
- Il brand "{data['brand']}" DEVE apparire nel meta title alla fine
- Il brand "{data['brand']}" DEVE apparire nella meta description
- Usa la capitalizzazione naturale italiana (solo prima parola maiuscola, nomi propri maiuscoli)
- Utilizza i dati REALI di volumi di ricerca e People Also Ask per ottimizzare il contenuto
- Considera le keyword correlate di SEMrush per arricchire il contenuto
- Analizza il featured snippet esistente per superarlo

Genera un content brief completo che includa:

1. **ANALISI INTENTO DI RICERCA E OBIETTIVO BASATA SU DATI REALI**
   - Analizza l'intento basandoti sui dati di volume SEMrush e le PAA di Google
   - Definisci l'obiettivo considerando la competizione reale (CPC e competition score)
   - Strategia per battere i competitor analizzando featured snippet e top 10

2. **LINEE GUIDA SEO E STILE CON KEYWORD DATA-DRIVEN**
   - Utilizza keyword correlate di SEMrush per densità e posizionamento
   - Incorpora naturalmente le ricerche correlate di Google
   - Stile basato sui volumi di ricerca e competizione
   - Lunghezza ottimale basata sull'analisi competitor

3. **META TITLE E DESCRIPTION OTTIMIZZATE**
   - Meta title (50-60 caratteri) con keyword principale e brand
   - Meta description (150-160 caratteri) che incorpori PAA e brand
   - Capitalizzazione naturale italiana

4. **STRATEGIA E-E-A-T POTENZIATA CON DATI REALI**
   
   **EXPERIENCE:** 
   - Esempi basati su ricerche correlate reali di Google
   - Casi studio che rispondano alle PAA
   
   **EXPERTISE:**
   - Dati statistici reali (volume ricerche, trend del settore)
   - Insight basati sui gap nei competitor
   
   **AUTHORITATIVENESS:**
   - Fonti che i competitor non citano
   - Opportunità di superare il featured snippet attuale
   
   **TRUSTWORTHINESS:**
   - Trasparenza su dati e fonti reali
   - Riconoscimento dei limiti dove competitors falliscono

5. **STRUTTURA CONTENUTO OTTIMIZZATA PER GOOGLE**
   - H1 ottimizzato per keyword principale
   - H2/H3 che rispondono alle People Also Ask
   - Sezioni per keyword correlate ad alto volume
   - Strategia per conquistare featured snippet
   - Risposte specifiche alle PAA di Google nel contenuto

6. **INTEGRAZIONE PEOPLE ALSO ASK**
   Per ogni PAA di Google:
   - Dove inserirla nella struttura
   - Come rispondere meglio dei competitor
   - Keyword correlate da usare nella risposta

7. **KEYWORD CORRELATE E SEMANTICHE**
   - Come integrare keyword correlate SEMrush
   - Ricerche correlate Google da includere
   - Densità ottimale basata su competition score

8. **LINK INTERNI STRATEGICI DATA-DRIVEN**
   - Link basati su analisi del traffico potenziale
   - Anchor text con keyword correlate
   - Distribuzione strategica per topic clustering

9. **FONTI E CREDIBILITÀ SUPERIORI AI COMPETITOR**
   - Fonti che mancano nei competitor top 10
   - Dati più recenti di quelli usati dai competitor
   - Authority sources che aumentano E-E-A-T

10. **CALL TO ACTION OTTIMIZZATE PER CONVERSIONE**
    - CTA basate sull'intento di ricerca reale
    - Micro-conversioni intermediate per high-competition keyword
    - CTA finali ottimizzate per volume/valore keyword

Utilizza ESCLUSIVAMENTE i dati reali forniti (volumi SEMrush, PAA Google, ricerche correlate) per creare una strategia superiore ai competitor. Focus su superare il featured snippet attuale e rispondere meglio alle PAA.
"""

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "Sei un esperto SEO content strategist che utilizza dati reali di SEMrush e Serper per creare content brief superiori alla concorrenza, ottimizzati per E-E-A-T e ranking Google."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=4000,
                temperature=0.7
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            st.error(f"Errore nella generazione del content brief: {str(e)}")
            return "Errore nella generazione del contenuto."

def create_docx(content: str, brand: str, topic: str) -> io.BytesIO:
    """Crea un documento DOCX formattato con il content brief"""
    doc = Document()
    
    # Configurazione del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Titolo principale
    title = doc.add_heading(f'Content brief SEO data-driven - {topic}', 0)
    title_format = title.runs[0].font
    title_format.name = 'Figtree'
    title_format.size = Pt(20)
    title_format.color.rgb = None
    
    # Sottotitolo
    subtitle = doc.add_paragraph(f'Brand: {brand}')
    subtitle_format = subtitle.runs[0].font
    subtitle_format.name = 'Figtree'
    subtitle_format.size = Pt(12)
    subtitle_format.bold = True
    
    doc.add_paragraph("")
    
    # Processa il contenuto
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], 1)
            heading_format = heading.runs[0].font
            heading_format.name = 'Figtree'
            heading_format.size = Pt(17)
            
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], 2)
            heading_format = heading.runs[0].font
            heading_format.name = 'Figtree'
            heading_format.size = Pt(17)
            
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], 3)
            heading_format = heading.runs[0].font
            heading_format.name = 'Figtree'
            heading_format.size = Pt(17)
            
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.font.name = 'Figtree'
            run.font.size = Pt(11)
            run.font.bold = True
            
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.runs[0].font.name = 'Figtree'
            p.runs[0].font.size = Pt(11)
            
        else:
            if line:
                p = doc.add_paragraph(line)
                p.runs[0].font.name = 'Figtree'
                p.runs[0].font.size = Pt(11)
    
    # Salva in BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

def main():
    st.markdown('<h1 class="main-header">📝 Content Brief Generator Pro</h1>', unsafe_allow_html=True)
    st.markdown('<p style="text-align: center; color: #666;">Genera content brief con dati SEO reali da SEMrush e Serper</p>', unsafe_allow_html=True)
    
    # Sidebar per API Keys
    with st.sidebar:
        st.markdown("### 🔑 Configurazione API")
        
        # OpenAI API Key
        openai_api_key = st.text_input("OpenAI API Key", type="password", help="Inserisci la tua API key di OpenAI")
        
        # SEMrush API Key
        st.markdown("---")
        st.markdown("#### 📊 SEMrush (Opzionale)")
        semrush_api_key = st.text_input("SEMrush API Key", type="password", help="Per dati keyword reali, volumi di ricerca e keyword correlate")
        if semrush_api_key:
            st.success("✅ SEMrush configurato")
        else:
            st.info("💡 Aggiungi per dati keyword reali")
        
        # Serper API Key
        st.markdown("#### 🔍 Serper (Opzionale)")
        serper_api_key = st.text_input("Serper API Key", type="password", help="Per People Also Ask e analisi SERP in tempo reale")
        if serper_api_key:
            st.success("✅ Serper configurato")
        else:
            st.info("💡 Aggiungi per PAA da Google")
        
        st.markdown("---")
        st.markdown("### 🚀 Miglioramenti")
        st.markdown("""
        **Con SEMrush + Serper:**
        - 📈 Volumi di ricerca reali
        - 🎯 Keyword correlate precise
        - ❓ People Also Ask live
        - 🏆 Featured snippet analysis
        - 📊 Competition data
        """)
    
    if not openai_api_key:
        st.info("👈 Inserisci almeno la OpenAI API Key nella sidebar per iniziare")
        return
    
    # Inizializza i servizi
    seo_enhancer = SEODataEnhancer(semrush_api_key, serper_api_key)
    generator = ContentBriefGenerator(openai_api_key, seo_enhancer)
    
    # Form principale
    with st.form("content_brief_form"):
        st.markdown('<h2 class="section-header">📋 Informazioni cliente</h2>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            brand = st.text_input("🏢 Nome del brand", placeholder="Es: TassoMutuo")
            website = st.text_input("🌐 URL del sito", placeholder="https://www.esempio.it")
            sitemap_url = st.text_input("🗺️ URL Sitemap.xml", placeholder="https://www.esempio.it/sitemap.xml")
            topic = st.text_area("📝 Argomento del contenuto", placeholder="Descrivi l'argomento principale dell'articolo")
        
        with col2:
            keywords = st.text_area("🔍 Keyword utili", placeholder="keyword principale, keyword secondaria, keyword long tail")
            faqs = st.text_area("❓ Domande frequenti (PAA)", placeholder="Inserisci domande conosciute, verranno integrate con PAA reali da Google se Serper è configurato")
            
            # Tone of voice con multiselect
            tone_options = [
                "Professionale", "Minimalista", "Persuasivo", 
                "Informativo", "Ricercato", "Popolare", "Personalizzato"
            ]
            tone_of_voice = st.multiselect("🎯 Tone of voice", tone_options, default=["Professionale"])
        
        # Preview API status
        if semrush_api_key or serper_api_key:
            st.markdown('<div class="success-box">🎯 <strong>Modalità Enhanced SEO attiva!</strong> Il content brief includerà dati reali da:', unsafe_allow_html=True)
            if semrush_api_key:
                st.markdown("✅ SEMrush: volumi di ricerca, keyword correlate, competition data")
            if serper_api_key:
                st.markdown("✅ Serper: People Also Ask live, ricerche correlate, featured snippets")
            st.markdown('</div>', unsafe_allow_html=True)
        
        # Gestione sitemap alternativa
        st.markdown('<h3 class="section-header">🔗 URL interne (fallback)</h3>', unsafe_allow_html=True)
        st.markdown('<div class="info-box">Se la sitemap non è accessibile, puoi incollare manualmente le URL interne principali (una per riga)</div>', unsafe_allow_html=True)
        
        manual_urls = st.text_area("📎 URL interne manuali", placeholder="https://www.sito.it/pagina1\nhttps://www.sito.it/pagina2", height=100)
        
        st.markdown('<h3 class="section-header">🔍 Analisi competitor</h3>', unsafe_allow_html=True)
        st.markdown('<div class="info-box">🎯 <strong>Inserimento manuale obbligatorio</strong>: Incolla il contenuto testuale completo di ogni competitor per un\'analisi più precisa e affidabile</div>', unsafe_allow_html=True)
        
        competitor_data = []
        for i in range(3):
            st.markdown(f"**📊 Competitor {i+1}**")
            
            # URL del competitor (opzionale, solo per riferimento)
            url = st.text_input(f"🌐 URL Competitor {i+1} (opzionale)", key=f"comp_url_{i}", placeholder="https://competitor.example.com/articolo", help="Solo per riferimento nel brief")
            
            # Contenuto testuale (obbligatorio)
            manual_content = st.text_area(
                f"📝 Contenuto testuale completo Competitor {i+1} *", 
                key=f"comp_content_{i}", 
                placeholder="Incolla qui tutto il contenuto testuale della pagina competitor (titoli, paragrafi, liste, etc.)",
                height=200,
                help="Copia e incolla tutto il testo della pagina competitor per un'analisi completa"
            )
            
            # Input aggiuntivi per analisi più dettagliata
            col1, col2 = st.columns(2)
            with col1:
                comp_title = st.text_input(f"📋 Title tag Competitor {i+1}", key=f"comp_title_{i}", placeholder="Title tag della pagina competitor")
            with col2:
                comp_meta = st.text_input(f"📄 Meta description Competitor {i+1}", key=f"comp_meta_{i}", placeholder="Meta description della pagina")
            
            if manual_content.strip():  # Solo se c'è contenuto
                competitor_data.append({
                    'url': url if url else f"Competitor {i+1}",
                    'manual_content': manual_content,
                    'manual_title': comp_title,
                    'manual_meta': comp_meta,
                    'competitor_number': i+1
                })
        
        # Mostra quanti competitor sono stati inseriti
        if competitor_data:
            st.success(f"✅ {len(competitor_data)} competitor pronti per l'analisi")
        
        submitted = st.form_submit_button("🚀 Genera content brief con dati SEO reali", use_container_width=True)
    
    if submitted:
        # Validazione input migliorata
        if not all([brand, website, topic, keywords]):
            st.error("❌ Compila tutti i campi obbligatori: Brand, Website, Argomento e Keywords")
            return
        
        if not competitor_data:
            st.error("❌ Inserisci il contenuto di almeno un competitor per procedere con l'analisi")
            return
        
        # Controllo che ci sia contenuto nei competitor
        valid_competitors_count = len([c for c in competitor_data if c['manual_content'].strip()])
        if valid_competitors_count == 0:
            st.error("❌ Inserisci il contenuto testuale di almeno un competitor")
            return
        
        with st.spinner("🔄 Generazione del content brief con dati SEO reali..."):
            progress_bar = st.progress(0)
            
            # Step 1: Analisi keyword con API avanzata
            if semrush_api_key or serper_api_key:
                st.info("🔍 Analisi keyword avanzata con SEMrush e Serper...")
                keyword_analysis = generator.analyze_keywords_with_apis(keywords)
                progress_bar.progress(15)
                
                # Mostra preview dati SEO avanzati
                if keyword_analysis['semrush_data'].get('status') == 'success':
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("📈 Volume ricerca", f"{keyword_analysis['semrush_data']['search_volume']:,}")
                    with col2:
                        st.metric("💰 CPC", f"€{keyword_analysis['semrush_data']['cpc']:.2f}")
                    with col3:
                        st.metric("⚔️ Competition", f"{keyword_analysis['semrush_data']['competition']:.2f}")
                    with col4:
                        intent = 'Transactional' if keyword_analysis['semrush_data']['cpc'] > 2 else 'Commercial' if keyword_analysis['semrush_data']['cpc'] > 1 else 'Informational'
                        st.metric("🎯 Intent", intent)
                
                # Mostra cluster tematici
                if keyword_analysis.get('topic_clusters'):
                    st.success(f"✅ Identificati {len(keyword_analysis['topic_clusters'])} cluster tematici")
                
                # Mostra analisi PAA
                if keyword_analysis.get('serper_data', {}).get('people_also_ask'):
                    paa_count = len(keyword_analysis['serper_data']['people_also_ask'])
                    st.success(f"✅ Analizzate {paa_count} People Also Ask con classificazione intent")
                
                # Mostra intent distribution
                if keyword_analysis.get('intent_categories'):
                    intent_summary = []
                    for intent, kws in keyword_analysis['intent_categories'].items():
                        if kws:
                            intent_summary.append(f"{intent}: {len(kws)} keyword")
                    if intent_summary:
                        st.info(f"📊 Distribuzione intent: {', '.join(intent_summary)}")
            else:
                st.info("📊 Analisi keyword base (senza API esterne)...")
                keyword_analysis = {
                    'main_keyword': keywords.split(',')[0].strip(), 
                    'semrush_data': {}, 
                    'serper_data': {},
                    'intent_categories': {},
                    'topic_clusters': {}
                }
                progress_bar.progress(15)
            
            # Step 2: Estrazione sitemap
            st.info("📡 Estrazione URL dalla sitemap...")
            sitemap_urls = []
            sitemap_error = False
            
            if sitemap_url:
                sitemap_urls = generator.get_sitemap_urls(sitemap_url)
                if not sitemap_urls:
                    sitemap_error = True
                    st.warning("⚠️ Impossibile accedere alla sitemap. Utilizzo URL manuali.")
            
            if sitemap_error or manual_urls.strip():
                manual_url_list = [url.strip() for url in manual_urls.split('\n') if url.strip()]
                sitemap_urls.extend(manual_url_list)
            
            progress_bar.progress(35)
            
            # Step 3: Preparazione dati competitor (solo manuali)
            st.info("📊 Elaborazione contenuti competitor inseriti manualmente...")
            competitors_processed = []
            
            for comp_data in competitor_data:
                if comp_data['manual_content'].strip():
                    # Processa il contenuto manuale
                    content = comp_data['manual_content']
                    word_count = len(content.split())
                    
                    # Estrae titoli dal contenuto se possibile
                    headings = []
                    lines = content.split('\n')
                    for line in lines:
                        line = line.strip()
                        # Cerca pattern di titoli
                        if (line.isupper() and len(line) > 10) or \
                           (line.startswith(('1.', '2.', '3.', '4.', '5.', '•', '-')) and len(line) > 15) or \
                           (len(line) < 100 and line.endswith((':', '?')) and len(line) > 10):
                            headings.append(line)
                    
                    processed_data = {
                        'url': comp_data['url'],
                        'title': comp_data['manual_title'] or f"Competitor {comp_data['competitor_number']}",
                        'meta_description': comp_data['manual_meta'],
                        'meta_keywords': "",
                        'headings': '\n'.join(headings[:15]),  # Primi 15 possibili titoli
                        'content': content,
                        'paragraphs': content.split('\n\n')[:10],  # Primi 10 paragrafi
                        'lists': [],
                        'word_count': word_count,
                        'status': 'manual',
                        'competitor_number': comp_data['competitor_number']
                    }
                    competitors_processed.append(processed_data)
            
            valid_competitors = competitors_processed
            progress_bar.progress(75)
            
            # Step 4: Preparazione dati
            st.info("📊 Preparazione dati per analisi E-E-A-T...")
            data = {
                'brand': brand,
                'website': website,
                'topic': topic,
                'keywords': keywords,
                'faqs': faqs,
                'tone_of_voice': tone_of_voice,
                'competitors': valid_competitors,
                'sitemap_urls': sitemap_urls,
                'manual_urls': manual_urls
            }
            progress_bar.progress(85)
            
            # Step 5: Generazione content brief
            st.info("🤖 Generazione content brief con AI e dati SEO reali...")
            content_brief = generator.generate_content_brief(data, keyword_analysis)
            progress_bar.progress(95)
            
            # Step 6: Creazione documento
            st.info("📄 Creazione documento DOCX...")
            docx_buffer = create_docx(content_brief, brand, topic)
            progress_bar.progress(100)
            
            st.success("✅ Content brief SEO data-driven generato con successo!")
        
        # Mostra risultati
        st.markdown('<h2 class="section-header">📋 Content brief SEO data-driven generato</h2>', unsafe_allow_html=True)
        
        # Dati SEO summary avanzati
        if keyword_analysis.get('semrush_data') or keyword_analysis.get('serper_data'):
            st.markdown('<h3 class="section-header">📊 Dati SEO reali utilizzati - Analisi avanzata</h3>', unsafe_allow_html=True)
            
            col1, col2 = st.columns(2)
            
            with col1:
                if keyword_analysis.get('semrush_data', {}).get('status') == 'success':
                    st.markdown("**🎯 Dati SEMrush**")
                    semrush = keyword_analysis['semrush_data']
                    st.write(f"• Volume: {semrush['search_volume']:,} ricerche/mese")
                    st.write(f"• CPC: €{semrush['cpc']:.2f}")
                    st.write(f"• Competition: {semrush['competition']:.2f}/1.0")
                    
                    intent_level = 'Transactional' if semrush['cpc'] > 2 else 'Commercial' if semrush['cpc'] > 1 else 'Informational'
                    st.write(f"• Intent primario: {intent_level}")
                    
                    if keyword_analysis.get('related_keywords'):
                        st.write(f"• {len(keyword_analysis['related_keywords'])} keyword correlate")
                    
                    if keyword_analysis.get('topic_clusters'):
                        st.write(f"• {len(keyword_analysis['topic_clusters'])} cluster tematici")
            
            with col2:
                if keyword_analysis.get('serper_data', {}).get('status') == 'success':
                    st.markdown("**🔍 Dati Serper**")
                    serper = keyword_analysis['serper_data']
                    if serper.get('people_also_ask'):
                        st.write(f"• {len(serper['people_also_ask'])} People Also Ask")
                    if serper.get('related_searches'):
                        st.write(f"• {len(serper['related_searches'])} ricerche correlate")
                    if serper.get('featured_snippet'):
                        snippet_analysis = serper.get('snippet_analysis', {})
                        st.write("• Featured snippet analizzato:")
                        st.write(f"  - Tipo: {snippet_analysis.get('structure_type', 'paragraph')}")
                        st.write(f"  - Lunghezza: {snippet_analysis.get('word_count', 0)} parole")
                    
                    # Mostra distribuzione intent PAA
                    if serper.get('paa_intents'):
                        paa_intents = serper['paa_intents']
                        intent_counts = {k: len(v) for k, v in paa_intents.items() if v}
                        if intent_counts:
                            st.write(f"• PAA per intent: {intent_counts}")
            
            # Preview cluster tematici
            if keyword_analysis.get('topic_clusters'):
                with st.expander("🎯 Cluster tematici identificati"):
                    sorted_clusters = sorted(keyword_analysis['topic_clusters'].items(), 
                                           key=lambda x: x[1]['total_volume'], reverse=True)
                    for cluster_name, cluster_data in sorted_clusters[:5]:
                        st.write(f"**{cluster_name.title()}**: {cluster_data['total_volume']:,} volume totale, {len(cluster_data['keywords'])} keyword")
            
            # Preview intent distribution
            if keyword_analysis.get('intent_categories'):
                with st.expander("🎭 Distribuzione keyword per intent"):
                    for intent, keywords in keyword_analysis['intent_categories'].items():
                        if keywords:
                            st.write(f"**{intent.title()}** ({len(keywords)} keyword):")
                            top_kws = sorted(keywords, key=lambda x: x['search_volume'], reverse=True)[:3]
                            for kw in top_kws:
                                st.write(f"  • {kw['keyword']} ({kw['search_volume']:,} vol.)")
            
            # Preview PAA classificate
            if keyword_analysis.get('serper_data', {}).get('paa_intents'):
                with st.expander("❓ People Also Ask classificate per intent"):
                    paa_intents = keyword_analysis['serper_data']['paa_intents']
                    for intent, questions in paa_intents.items():
                        if questions:
                            st.write(f"**{intent.title()}**:")
                            for q in questions[:3]:
                                st.write(f"  • {q}")
        
        # Preview del contenuto
        with st.expander("👁️ Anteprima content brief", expanded=True):
            st.markdown(content_brief)
        
        # Download button
        st.download_button(
            label="📥 Scarica content brief SEO data-driven (DOCX)",
            data=docx_buffer.getvalue(),
            file_name=f"content_brief_SEO_{brand}_{topic.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
        # Informazioni aggiuntive
        st.markdown('<h3 class="section-header">📊 Riepilogo analisi completa</h3>', unsafe_allow_html=True)
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("🔍 Competitor analizzati", len(valid_competitors))
        with col2:
            st.metric("🔗 URL interne trovate", len(sitemap_urls))
        with col3:
            paa_count = len(keyword_analysis.get('serper_data', {}).get('people_also_ask', []))
            st.metric("❓ PAA analizzate", paa_count)
        with col4:
            cluster_count = len(keyword_analysis.get('topic_clusters', {}))
            st.metric("🎯 Topic cluster", cluster_count)
        
        # Metriche aggiuntive
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            kw_count = len(keyword_analysis.get('related_keywords', []))
            st.metric("📈 Keyword correlate", kw_count)
        with col2:
            total_volume = sum(kw['search_volume'] for kw in keyword_analysis.get('related_keywords', []))
            st.metric("📊 Volume totale", f"{total_volume:,}")
        with col3:
            avg_competition = sum(kw['competition'] for kw in keyword_analysis.get('related_keywords', [])) / len(keyword_analysis.get('related_keywords', [])) if keyword_analysis.get('related_keywords') else 0
            st.metric("⚔️ Competition media", f"{avg_competition:.2f}")
        with col4:
            intent_types = len([k for k, v in keyword_analysis.get('intent_categories', {}).items() if v])
            st.metric("🎭 Intent types", intent_types)
        
        # Dettagli competitor
        with st.expander("🔍 Dettagli analisi competitor"):
            for comp in valid_competitors:
                st.markdown(f"**Competitor {comp['competitor_number']}:** {comp['url']}")
                st.write(f"📝 Contenuto inserito manualmente")
                if comp['title'] != f"Competitor {comp['competitor_number']}":
                    st.write(f"Titolo: {comp['title']}")
                if comp['meta_description']:
                    st.write(f"Meta Description: {comp['meta_description']}")
                st.write(f"Numero parole: {comp['word_count']}")
                if comp['headings']:
                    st.write(f"Strutture trovate: {len(comp['headings'].split())}")
                st.markdown("---")
        
        # Enhanced E-E-A-T checklist
        st.markdown('<h3 class="section-header">🏆 Checklist E-E-A-T Enhanced</h3>', unsafe_allow_html=True)
        st.markdown('<div class="success-box">Il content brief generato include suggerimenti basati su <strong>dati SEO reali</strong> per massimizzare E-E-A-T:</div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("""
            **📚 Experience (Data-driven)**
            - ✅ Esempi basati su ricerche correlate reali
            - ✅ Casi studio che rispondono alle PAA
            - ✅ Scenari ottimizzati per volumi di ricerca
            - ✅ Testimonianze mirate al search intent
            
            **🎓 Expertise (SEO-optimized)**
            - ✅ Insight basati su keyword correlate
            - ✅ Dati statistici reali (volumi, trend)
            - ✅ Gap analysis vs competitor top 10
            - ✅ Featured snippet optimization
            """)
        
        with col2:
            st.markdown("""
            **⭐ Authoritativeness (SERP-focused)**
            - ✅ Fonti che battono competitor attuali
            - ✅ Strategia per featured snippet
            - ✅ Authority building per high-volume keywords
            - ✅ Topical relevance enhancement
            
            **🛡️ Trustworthiness (Search-aligned)**
            - ✅ Trasparenza su dati e metodologie
            - ✅ Riconoscimento limiti vs competitor
            - ✅ User intent satisfaction
            - ✅ Search quality guidelines compliance
            """)
        
        # API usage tips
        if not semrush_api_key or not serper_api_key:
            st.markdown('<h3 class="section-header">💡 Migliora ulteriormente i risultati</h3>', unsafe_allow_html=True)
            st.markdown('<div class="info-box">', unsafe_allow_html=True)
            if not semrush_api_key:
                st.write("🔹 **Aggiungi SEMrush API** per: volumi di ricerca precisi, keyword difficulty, competitor traffic analysis")
            if not serper_api_key:
                st.write("🔹 **Aggiungi Serper API** per: People Also Ask live, featured snippets, real-time SERP data")
            st.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
